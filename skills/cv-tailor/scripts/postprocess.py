"""
postprocess.py — post-render pass over a rendered CV .docx.

Why this exists
---------------
docxtpl RichText passed through a plain `{{ bullet }}` placeholder embeds the
RichText's run-XML *inside* the placeholder's <w:t> element. That is invalid
OOXML: Word, python-docx, and ATS parsers all read the paragraph as EMPTY.
This was the 2026-05-11 incident and the corruption that shipped in every
labeled-mode CV of the 2026-06-25 (Cairo) and 2026-06-27 (Berlin) batches.

So RichText is banned from the render path. Bullets render as plain strings
(valid OOXML by construction, formatting inherited from the styled placeholder
run), and bold is applied HERE, after `tpl.save()`:

  1. `build_bold_plan()` (md_to_richtext.py) strips the `**` markers before
     render and records which character spans of each bullet should be bold.
  2. `postprocess_cv()` re-opens the saved .docx with python-docx, finds each
     bullet paragraph, and splits its single run at the span boundaries by
     CLONING the run element. The clone inherits the template run's full rPr
     (Calibri, sz 20, paired w:b/w:bCs) with zero re-declaration in code —
     the formatting constants in docxtpl-recipe.md stay template-owned.

The same pass removes disabled sections (config `cv.sections` /
`cv.region_section_overrides`, e.g. EU: summary: false) by deleting the
section's header paragraph through the paragraph before the next header.

Every bullet in the plan is located by exact (whitespace-normalized) text
match with a forward-moving cursor. A bullet that cannot be found raises
PostprocessError — a loud render-time failure instead of a silent blank CV.
"""

import copy
import urllib.parse

from docx import Document
from docx.oxml.ns import qn
from docx.text.run import Run


# Literal header-paragraph text per section in the OPUS template. A different
# template supplies its own mapping via the `section_headers` argument.
OPUS_SECTION_HEADERS = {
    "summary": "PROFESSIONAL SUMMARY",
    "core_skills": "CORE SKILLS",
    "experience": "PROFESSIONAL EXPERIENCE",
    "education": "EDUCATION",
    "additional": "ADDITIONAL",
}

# Sections that must never be removed, whatever the config says.
REQUIRED_SECTIONS = ("tagline", "contact", "experience", "education")

# Content-map keys whose rendered text sits inside a <w:hyperlink>. Each
# one's relationship target must equal its own visible text — see
# _sync_contact_hyperlinks.
CONTACT_LINK_KEYS = ("personal_site", "linkedin_url")


class PostprocessError(ValueError):
    """A bold-plan entry could not be matched to the rendered document."""


class ContactLinkError(ValueError):
    """A contact hyperlink could not be uniquely bound to its target."""


def _norm(text):
    return " ".join(text.split())


def _find_paragraph(paragraphs, text, start=0):
    """Index of the first paragraph (from `start`) whose normalized text
    equals `text` normalized, or None."""
    target = _norm(text)
    for i in range(start, len(paragraphs)):
        if _norm(paragraphs[i].text) == target:
            return i
    return None


def _merge_runs(paragraph):
    """Collapse the paragraph to a single run carrying all its text.

    docxtpl renders a `{{ placeholder }}` inside one run, so bullets arrive
    as a single run; this is defensive for a paragraph that arrives split
    (proofing marks, rsid edits in a future template).
    """
    runs = paragraph.runs
    if len(runs) <= 1:
        return
    runs[0].text = "".join(r.text for r in runs)
    for r in runs[1:]:
        r._r.getparent().remove(r._r)


def _segments(text, spans):
    """Split `text` into ordered (segment_text, bold) pairs from bold spans.

    Spans are clamped to the text, empties dropped, overlapping/adjacent
    spans merged.
    """
    cleaned = []
    for s, e in sorted((max(0, s), min(len(text), e)) for s, e in spans):
        if e <= s:
            continue
        if cleaned and s <= cleaned[-1][1]:
            cleaned[-1] = (cleaned[-1][0], max(cleaned[-1][1], e))
        else:
            cleaned.append((s, e))
    segments = []
    pos = 0
    for s, e in cleaned:
        if s > pos:
            segments.append((text[pos:s], False))
        segments.append((text[s:e], True))
        pos = e
    if pos < len(text):
        segments.append((text[pos:], False))
    return [seg for seg in segments if seg[0]]


def _split_run_bold(paragraph, spans):
    """Rebuild the paragraph's run as bold/plain segments per `spans`.

    Each segment is a deep copy of the original run element, so the
    template's rPr (font, size, paired w:b/w:bCs attributes) is inherited
    exactly; only the bold flag differs. Returns the number of bold runs
    created.
    """
    _merge_runs(paragraph)
    if not paragraph.runs:
        raise PostprocessError(
            "bullet paragraph has no runs: " + repr(paragraph.text[:60]))
    src = paragraph.runs[0]
    segments = _segments(src.text, spans)
    if not any(bold for _, bold in segments):
        return 0

    prev_el = src._r
    bolded = 0
    for seg_text, seg_bold in segments:
        el = copy.deepcopy(src._r)
        prev_el.addnext(el)
        prev_el = el
        run = Run(el, paragraph)
        run.text = seg_text  # python-docx sets xml:space="preserve" as needed
        if seg_bold:
            run.font.bold = True
            run.font.cs_bold = True  # keep w:b / w:bCs paired (template rule)
            bolded += 1
    src._r.getparent().remove(src._r)
    return bolded


def _apply_bold(doc, bold_plan):
    """Locate every planned bullet in document order; bold its spans.

    Entries with no spans are still located — so this doubles as a
    render-time integrity assert: a bullet the plan expected but the
    document cannot show is a hard failure, not a silent blank.
    """
    if not bold_plan:
        return 0
    paras = doc.paragraphs
    cursor = 0
    bolded = 0
    for spec in bold_plan:
        idx = _find_paragraph(paras, spec["text"], start=cursor)
        if idx is None:
            raise PostprocessError(
                "bullet not found in rendered doc (order or text mismatch): "
                + repr(spec["text"][:60]))
        cursor = idx + 1
        if spec.get("spans"):
            bolded += _split_run_bold(paras[idx], spec["spans"])
    return bolded


def _normalize_target(value):
    """Return an http(s) target for a contact value.

    A bare `example.com` gets an https:// scheme — Word needs an explicit one
    or the relationship is not a working external link. Anything that is not
    http/https is rejected rather than normalized: a contact line is a trust
    boundary, and `javascript:` or `file:` targets have no business in a CV
    that gets mailed to strangers.
    """
    value = (value or "").strip()
    if not value:
        raise ContactLinkError("empty contact link value")
    scheme = urllib.parse.urlparse(value).scheme.lower()
    if not scheme:
        return "https://" + value
    if scheme not in ("http", "https"):
        raise ContactLinkError(
            "unsupported contact link scheme %r in %r (http/https only)"
            % (scheme, value))
    return value


def _sync_contact_hyperlinks(doc, contact_links):
    """Point each contact hyperlink's relationship at its own visible text.

    Why this exists: docxtpl renders `{{ personal_site }}` into the <w:t>
    inside a <w:hyperlink>, but a hyperlink's DESTINATION lives in
    word/_rels/document.xml.rels, keyed by r:id — a part docxtpl never
    touches. A template built from someone's finished CV therefore keeps
    THEIR site and LinkedIn as the click targets of every CV rendered from
    it forever, visible to any ATS that reads relationships instead of
    display text. That is the 2026-08 leak.

    Each link is bound by r:id and its target set from the content map. The
    old target is never consulted — matching on it would re-couple the fix
    to the very values being removed, and would break the moment the
    template is neutralized.
    """
    rels = doc.part.rels
    elements = list(doc.element.iter(qn("w:hyperlink")))
    synced = 0
    for key in CONTACT_LINK_KEYS:
        if key not in contact_links:
            continue
        raw = contact_links[key]
        label = _norm(raw)
        target = _normalize_target(raw)
        matches = []
        for el in elements:
            rid = el.get(qn("r:id"))
            if rid is None or rid not in rels:
                continue
            text = _norm("".join(n.text or "" for n in el.iter(qn("w:t"))))
            if text == label:
                matches.append(rid)
        if len(matches) != 1:
            raise ContactLinkError(
                "%s (%r): expected exactly 1 hyperlink with that visible "
                "text, found %d — cannot bind its target safely"
                % (key, label, len(matches)))
        rels[matches[0]]._target = target
        synced += 1
    return synced


def _remove_sections(doc, disabled, headers):
    """Delete each disabled section's paragraphs (header through the
    paragraph before the next known header). Collect-then-remove."""
    removed, warnings = [], []
    if not disabled:
        return removed, warnings

    paras = doc.paragraphs
    header_positions = {}
    for i, p in enumerate(paras):
        stripped = p.text.strip()
        for name, header_text in headers.items():
            if stripped == header_text:
                header_positions[name] = i
    header_indices = sorted(header_positions.values())

    ranges = []
    for name in disabled:
        if name in REQUIRED_SECTIONS:
            warnings.append(
                "section '%s' is required and was not removed" % name)
            continue
        if name not in headers:
            warnings.append(
                "section '%s' has no header mapping; not removed" % name)
            continue
        if name not in header_positions:
            warnings.append(
                "header for section '%s' not found in document; "
                "nothing removed" % name)
            continue
        start = header_positions[name]
        following = [i for i in header_indices if i > start]
        end = following[0] if following else len(paras)
        ranges.append((name, start, end))

    for name, start, end in ranges:
        for p in paras[start:end]:
            el = p._element
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        removed.append(name)
    return removed, warnings


def _move_section_before(doc, move_name, anchor_name, headers):
    """Move `move_name`'s paragraphs (header through the paragraph before
    the next known header) to just before `anchor_name`'s header.

    Used by student mode to put EDUCATION above PROFESSIONAL EXPERIENCE.
    Returns (moved, warnings); no-op when either header is missing.
    """
    paras = doc.paragraphs
    header_positions = {}
    for i, p in enumerate(paras):
        stripped = p.text.strip()
        for name, header_text in headers.items():
            if stripped == header_text:
                header_positions[name] = i

    if move_name not in header_positions or anchor_name not in header_positions:
        missing = [n for n in (move_name, anchor_name)
                   if n not in header_positions]
        return False, ["header for section '%s' not found; section not moved"
                       % "', '".join(missing)]
    start = header_positions[move_name]
    if header_positions[anchor_name] >= start:
        return False, []  # already above the anchor
    following = [i for i in sorted(header_positions.values()) if i > start]
    end = following[0] if following else len(paras)
    anchor_el = paras[header_positions[anchor_name]]._element
    for p in paras[start:end]:
        anchor_el.addprevious(p._element)
    return True, []


def postprocess_cv(docx_path, bold_plan, disabled_sections=(),
                   section_headers=OPUS_SECTION_HEADERS, student_mode=False,
                   contact_links=None):
    """Post-render pass: remove disabled sections, apply planned bold.

    `student_mode=True` moves EDUCATION above PROFESSIONAL EXPERIENCE
    (early-career CVs lead with education).

    Opens the document once, saves in place. Returns a summary dict:
    {"removed_sections": [...], "bolded_runs": int, "warnings": [...]}.

    `contact_links` is the {key: value} slice of the content map whose text
    renders inside a hyperlink (CONTACT_LINK_KEYS). When given, each link's
    relationship target is rebound to its own visible value; see
    _sync_contact_hyperlinks. Callers that omit it leave the template's
    targets in place — audit check 5 is the independent backstop.

    Raises PostprocessError when a planned bullet cannot be located —
    ship nothing rather than a CV whose bullets are not where the plan
    says they are. Raises ContactLinkError when a contact hyperlink cannot
    be uniquely bound.
    """
    doc = Document(docx_path)
    removed, warnings = _remove_sections(doc, disabled_sections,
                                         section_headers)
    if student_mode:
        # Must run before _apply_bold: its forward cursor over
        # doc.paragraphs assumes final document order — so when the move
        # happens, the plan is reordered too (degree bullets first).
        moved, move_warnings = _move_section_before(
            doc, "education", "experience", section_headers)
        warnings.extend(move_warnings)
        if moved:
            bold_plan = ([s for s in bold_plan if s["section"] == "degree"]
                         + [s for s in bold_plan if s["section"] != "degree"])
    bolded = _apply_bold(doc, bold_plan)
    synced = _sync_contact_hyperlinks(doc, contact_links or {})
    doc.save(docx_path)
    return {
        "removed_sections": removed,
        "bolded_runs": bolded,
        "synced_contact_links": synced,
        "warnings": warnings,
    }


if __name__ == "__main__":
    # Smallest self-check: segments logic (the only branchy pure function).
    assert _segments("abcdef", [(1, 3)]) == [("a", False), ("bc", True),
                                             ("def", False)]
    assert _segments("abcdef", [(0, 2), (2, 4)]) == [("abcd", True),
                                                     ("ef", False)]
    assert _segments("abc", [(5, 9)]) == [("abc", False)]
    assert _segments("Label: rest", [(0, 6)]) == [("Label:", True),
                                                  (" rest", False)]
    print("postprocess self-test (_segments): passed.")
