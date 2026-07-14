"""
audit.py — post-render audit for tailored CVs.

Implements the programmatic checks of the post-render audit
(see skills/cv-tailor/references/post-render-audit.md). The editorial checks (#1, #3)
are run by the model, not here.

Programmatic checks:
  Check 2 — at least two JD keywords appear verbatim in the rendered CV.
  Check 4 — every & from the content_map survives into the rendered XML.
  Check 5 — rendered-text integrity: every authored bullet is readable via
            python-docx (the parser Word and ATS systems agree with), per-slot
            counts match, no raw markup in any paragraph text, planned bold
            spans are real bold runs, and the contact hyperlinks survive.
            Supersedes the old bold-run regex count, which could never fail
            on the OPUS template (its section headers are bold) and passed
            the 2026-05-11 / test5 corruption where every bullet rendered
            EMPTY (RichText embedded inside <w:t> is invisible to Word but
            visible to a regex).
  Check 6 — no em dashes in the rendered CV (employer-facing output; see
            shared/conventions.md).
  Check 7 — experience section is in strict reverse-chronological order and
            the primary employer's contiguous block occupies slots 1 + 2.
  Check 8 — tailoring coverage: every experience slot reflects the diagnosis
            (each slot's bullets carry at least one diagnosed keyword), so no
            slot ships as un-angled career-file boilerplate.
  Check 9 — grounding: every number/percentage/count in the rendered CV traces
            to the career file, catching invented or inflated metrics.
  Check 10 — bullet strength / proof density: no experience bullet hides
            behind a generic abstraction *while carrying no concrete proof of
            its own*, and (when the career file is supplied) each slot clears
            a proof-density floor — most of its bullets must carry a digit or
            a named entity that actually appears in the career file. One
            interpretive/ungrounded bullet per 3-bullet slot is allowed by
            design (the domain-translation pattern).
  Check 11 — proof-point presence: the diagnosis's per-slot "proof point:"
            (parsed from Diagnosis.md — the source of truth, not a
            model-copied field) must surface in that slot's bullets. Catches
            "the 30% fell out of the lead slot".
  Check 12 — education completeness: the `degrees` list carries every degree
            (>= cv.expected_degree_count when configured) and each degree's
            institution is visible in the rendered document. Catches the
            2026-07-14 dropped-BA failure.

Editorial checks 1 and 3 are REQUIRED: run_full_audit seeds them as failed,
and all_passed stays False until the model records a verdict for each via
result.record_editorial(...). Recording them is authoring work, not a pause.

Batch-level (not part of run_full_audit): scan_batch_sameyness(session_dir)
warns on exact-duplicate bullets across different CVs in a session folder.

A CV that fails any check is NOT shipped.
"""

import os
import zipfile
import re
from dataclasses import dataclass, field

from docx import Document


# The two editorial checks the model must grade explicitly. Seeded as failed
# by run_full_audit so a CV can never pass by omission — the 2026-06-27
# Berlin batch shipped with these silently skipped.
EDITORIAL_CHECKS = ("check_1_lead_slots", "check_3_recruiter_fit")


@dataclass
class AuditResult:
    passed: dict = field(default_factory=dict)
    notes: dict = field(default_factory=dict)

    @property
    def all_passed(self):
        return all(self.passed.values())

    @property
    def failure_summary(self):
        lines = []
        for name, ok in self.passed.items():
            if not ok:
                lines.append("  FAIL [" + name + "]: " + self.notes.get(name, ""))
        return "\n".join(lines) if lines else "All checks passed."

    def record_editorial(self, check_name, ok, note):
        """Record the model's verdict for an editorial check (1 or 3).

        check_1_lead_slots — do the lead slots serve the diagnosed problem
        with their named proof points surfaced?
        check_3_recruiter_fit — richness vs the career file, domain
        translation, recruiter-fit, and the Check 9 honesty companion
        (no semantic inflation: "supported" did not become "led").
        """
        if check_name not in EDITORIAL_CHECKS:
            raise KeyError(
                "unknown editorial check %r; expected one of %s"
                % (check_name, EDITORIAL_CHECKS))
        self.passed[check_name] = bool(ok)
        self.notes[check_name] = note


def _read_document_xml(docx_path):
    """Return word/document.xml as a string."""
    with zipfile.ZipFile(docx_path) as z:
        return z.read("word/document.xml").decode("utf-8", errors="replace")


def _visible_text(document_xml):
    """Concatenate all <w:t> text content from the document XML."""
    return " ".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", document_xml, re.DOTALL))


def _bullet_text(bullet):
    """Return the plain text of a content_map bullet.

    Since v1.8.0 bullets are always plain strings (build_bold_plan strips
    the ** markers; RichText is banned from the render path). The RichText
    branch below is kept ONLY for forensics of files/content maps produced
    before v1.8.0 — do not rely on it for new renders.
    """
    if isinstance(bullet, str):
        return bullet
    xml = getattr(bullet, "xml", "") or str(bullet)
    return " ".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, re.DOTALL))


def _iter_strings(obj):
    """Yield every string value nested anywhere in obj."""
    if isinstance(obj, str):
        yield obj
    elif isinstance(obj, dict):
        for v in obj.values():
            yield from _iter_strings(v)
    elif isinstance(obj, (list, tuple)):
        for v in obj:
            yield from _iter_strings(v)


def check_2_keywords_in_experience(experiences, expected_keywords):
    """Check 2: >= 2 JD keywords appear verbatim in EXPERIENCE BULLETS.

    Reads the authored bullets; Check 5 (rendered-text integrity) guarantees
    they match the rendered document. The pre-v1.8.0 version searched the
    whole document's visible text, so a CV could pass on tagline/summary
    keywords while every experience bullet stayed generic — contradicting
    this check's own spec in post-render-audit.md.
    """
    if not expected_keywords:
        return True, "No keywords supplied; check skipped."
    text = " ".join(
        _bullet_text(b)
        for e in experiences
        for b in (e.get("bullets") or [])
    ).lower()
    hits = [kw for kw in expected_keywords if kw.lower() in text]
    ok = len(hits) >= 2
    if ok:
        note = str(len(hits)) + "/" + str(len(expected_keywords)) \
            + " keywords found verbatim in experience bullets: " + str(hits)
    else:
        note = ("Only " + str(len(hits)) + " keyword(s) found verbatim in "
                "experience bullets (" + str(hits) + "); need >= 2. Keywords "
                "in the tagline/summary/skills do not count for this check.")
    return ok, note


def check_4_ampersands(document_xml, content_map):
    """Check 4: every & in content_map values survives into the rendered XML.

    The pass/fail criterion is purely the ampersand count. Double spaces are
    NOT a fail criterion — many templates legitimately use spaced separators
    ("  |  ", "  .  "). A double space is only a locating aid once the count
    check has already failed.
    """
    expected_amp = 0
    for value in _iter_strings(content_map):
        expected_amp += value.count("&")
    rendered_amp = document_xml.count("&amp;")
    ok = rendered_amp >= expected_amp
    if ok:
        note = (str(rendered_amp) + " ampersands rendered (>= "
                + str(expected_amp) + " expected).")
    else:
        text = _visible_text(document_xml)
        double_space_hits = len(re.findall(r"\S  \S", text))
        note = ("expected >= " + str(expected_amp) + " ampersands, found "
                + str(rendered_amp) + " (autoescape=False suspected). "
                + str(double_space_hits) + " double-space occurrence(s) in "
                "visible text — inspect those to locate the stripped "
                "ampersand(s).")
    return ok, note


def _norm_ws(text):
    return " ".join(text.split())


def check_5_rendered_integrity(rendered_docx_path, content_map,
                               bold_plan=None, expect_bold=False):
    """Check 5: rendered-text integrity (supersedes the bold-run regex count).

    Opens the rendered docx with python-docx — the same parse Word and ATS
    systems perform — and asserts the authored content is actually visible:

      a. every authored experience bullet appears as a whole, non-empty
         paragraph, in order; per-slot readable counts match the content map;
      b. same for every degree's bullets (and the retired msc_bullets /
         ba_bullets keys when auditing a pre-v1.9.0 file);
      c. no paragraph text contains raw markup ('<w:') or leftover '**';
      d. when bold was planned: each planned span's text is covered by a run
         with run.bold True in its paragraph (real run inspection);
      e. the contact hyperlinks survived the postprocess round-trip
         (>= 2 hyperlink relationships — asherif.me + LinkedIn in OPUS).

    Why it exists: the old check counted <w:b/> in the raw XML. The OPUS
    template's own section headers are bold, so that count could never be
    zero — the check was unfailable — and it passed the 2026-05-11 / test5
    corruption where every bullet rendered EMPTY (RichText embedded inside
    <w:t> is invisible to Word but visible to a regex). This check reads
    what a recruiter reads.
    """
    problems = []
    doc = Document(rendered_docx_path)
    paras = doc.paragraphs

    def find(text, start):
        target = _norm_ws(text)
        for i in range(start, len(paras)):
            if _norm_ws(paras[i].text) == target:
                return i
        return None

    # a + b: every authored bullet is a readable paragraph, in order.
    expected_lists = []
    for si, role in enumerate(content_map.get("experiences", [])):
        label = "slot %d (%s)" % (si + 1, role.get("company", "?"))
        expected_lists.append((label, role.get("bullets", []) or []))
    for di, deg in enumerate(content_map.get("degrees", [])):
        if deg.get("bullets"):
            label = "degree %d (%s)" % (di + 1, deg.get("institution", "?"))
            expected_lists.append((label, deg["bullets"]))
    for key in ("msc_bullets", "ba_bullets"):  # pre-v1.9.0 forensics only
        if content_map.get(key):
            expected_lists.append((key, content_map[key]))

    cursor = 0
    total_expected = 0
    total_found = 0
    for label, bullets in expected_lists:
        found = 0
        for b in bullets:
            btext = _bullet_text(b)
            if not btext.strip():
                continue
            total_expected += 1
            idx = find(btext, cursor)
            if idx is None:
                problems.append(
                    "%s: bullet NOT readable in rendered doc: %r"
                    % (label, btext[:60]))
                continue
            cursor = idx + 1
            found += 1
            total_found += 1
        expected = len([b for b in bullets if _bullet_text(b).strip()])
        if found != expected:
            problems.append(
                "%s: only %d/%d bullets readable" % (label, found, expected))

    # c: no raw markup or leftover markers in any visible paragraph.
    for i, p in enumerate(paras):
        if "<w:" in p.text or "**" in p.text:
            problems.append(
                "paragraph %d contains raw markup or ** markers: %r"
                % (i, p.text[:60]))

    # d: planned bold spans rendered as real bold runs.
    bold_checked = 0
    if expect_bold and bold_plan:
        cursor2 = 0
        for spec in bold_plan:
            spans = spec.get("spans") or []
            if not spans:
                continue
            idx = find(spec["text"], cursor2)
            if idx is None:
                continue  # already reported unreadable above
            cursor2 = idx + 1
            bold_text = _norm_ws(
                "".join(r.text for r in paras[idx].runs if r.bold))
            for s, e in spans:
                span_text = _norm_ws(spec["text"][s:e])
                if span_text and span_text not in bold_text:
                    problems.append(
                        "bold span not rendered bold: %r in %r"
                        % (span_text[:40], spec["text"][:40]))
                else:
                    bold_checked += 1
    elif expect_bold and not bold_plan:
        problems.append("expect_bold=True but no bold_plan was supplied")

    # e: contact hyperlinks survived the round-trip.
    with zipfile.ZipFile(rendered_docx_path) as z:
        rels = z.read("word/_rels/document.xml.rels").decode(
            "utf-8", errors="replace")
    n_links = rels.count("/relationships/hyperlink")
    if n_links < 2:
        problems.append(
            "hyperlink relationships: %d found, >= 2 expected (personal "
            "site + LinkedIn) — the postprocess round-trip may have dropped "
            "them" % n_links)

    ok = not problems
    if ok:
        note = ("%d/%d bullets readable via python-docx; %d bold span(s) "
                "verified; %d hyperlink rel(s) intact."
                % (total_found, total_expected, bold_checked, n_links))
    else:
        note = "; ".join(problems)
    return ok, note


def check_6_no_em_dashes(document_xml):
    """Check 6: no em dashes in the rendered CV.

    Em dashes (U+2014, —) are banned from all employer-facing output.
    See shared/conventions.md.
    """
    text = _visible_text(document_xml)
    em_dash_count = text.count("—")
    ok = em_dash_count == 0
    if ok:
        note = "No em dashes found."
    else:
        note = (str(em_dash_count) + " em dash(es) found in rendered CV. "
                "Replace with commas, periods, or restructure the sentence. "
                "See shared/conventions.md.")
    return ok, note


def check_7_experience_structure(experiences):
    """Check 7: experience list is reverse-chronological and the primary
    employer's contiguous block holds slots 1 + 2.

    `experiences` is the list of dicts from the content_map (before rendering).
    Each dict must have: company, title, end_year (int; use 9999 for 'Present').
    An entry may set `concurrent: true` — an ongoing SIDE engagement (e.g.
    freelance) that overlaps the primary block; it is exempted from the
    reverse-chronology sort (it legitimately sits below the block) but not
    from the contiguous-block rule.

    Returns (ok, note). Returns (True, skip-note) only when experiences has
    fewer than 2 entries. A missing end_year is a FAIL, not a skip — the old
    skip-on-absence let the 2026-06-27 Berlin driver bypass this check
    entirely by passing end_years=None.
    """
    if not experiences or len(experiences) < 2:
        return True, "Fewer than 2 experience entries; structure check skipped."

    missing = [str(e.get("company", "slot " + str(i + 1)))
               for i, e in enumerate(experiences)
               if not isinstance(e.get("end_year"), int)]
    if missing:
        return False, (
            "end_year missing (or not an int) on: " + ", ".join(missing)
            + ". Required on every entry (9999 = Present) — chronology "
            "cannot be verified without it, and skip-on-absence was the "
            "dodge that let unordered CVs ship."
        )

    # Check strict reverse-chronological order (concurrent side roles exempt).
    ordered = [e for e in experiences if not e.get("concurrent")]
    years = [e["end_year"] for e in ordered]
    if years != sorted(years, reverse=True):
        return False, (
            "Experience entries are not in strict reverse-chronological order. "
            "Order: " + str([str(e.get("company", "?")) + " " + str(e.get("end_year"))
                             for e in experiences]) + ". "
            "Most recent role (highest end_year) must be slot 1. "
            "Ongoing roles use end_year=9999; an ongoing side engagement that "
            "belongs below the primary block must set concurrent: true."
        )

    # Check that slots 1 and 2 share the same employer (contiguous block rule)
    slot1_company = experiences[0].get("company", "")
    slot2_company = experiences[1].get("company", "") if len(experiences) > 1 else ""
    if slot1_company and slot2_company and slot1_company != slot2_company:
        return False, (
            "Slots 1 and 2 are different employers ("
            + str(slot1_company) + " vs " + str(slot2_company) + "). "
            "When the candidate has two adjacent roles at the same primary employer "
            "(e.g., Statista Expert + Statista Assistant), they must occupy slots "
            "1 + 2 as a contiguous block. See "
            "skills/cv-tailor/references/experience-slot-logic.md."
        )

    return True, ("Experience structure valid: reverse-chronological; "
                  "slots 1 + 2 share employer " + str(slot1_company) + ".")


def check_8_slot_coverage(experiences, expected_keywords):
    """Check 8: every experience slot is angled to this role.

    The diagnosis now mandates that at least one diagnosed keyword/angle reaches
    every experience slot, not just the lead (see diagnosis-template.md
    "Section angles"). The programmatic floor: each slot's bullets must contain
    at least one diagnosed keyword verbatim. A slot with zero is the symptom of
    un-angled career-file boilerplate pasted across CVs.

    Skipped (manual review) when keywords or bullets are absent.
    """
    if not expected_keywords or not experiences:
        return True, "No keywords or no experiences; coverage check skipped."
    if not all(e.get("bullets") for e in experiences):
        return True, ("Not all slots carry bullets in the content_map; "
                      "coverage check requires manual review.")

    kws = [k.lower() for k in expected_keywords]
    uncovered = []
    for i, e in enumerate(experiences):
        text = " ".join(_bullet_text(b) for b in e.get("bullets", [])).lower()
        if not any(k in text for k in kws):
            uncovered.append("slot " + str(i + 1) + " ("
                             + str(e.get("company", "?")) + ")")
    ok = not uncovered
    if ok:
        note = "Every experience slot carries >= 1 diagnosed keyword."
    else:
        note = ("Un-angled slot(s) with zero diagnosed keywords: "
                + ", ".join(uncovered) + ". The diagnosis must give each slot a "
                "Section-angle; do not paste career-file phrasing verbatim. See "
                "references/content-map-schema.md 'Facts vs angle'.")
    return ok, note


# Metric shapes Check 9 verifies against the career file. Deliberately
# excluded: bare integers, years, and letter-digit tokens (B2, Phase III) —
# flagging those would fail legitimate dates and language levels.
_METRIC_RE = re.compile(
    r"\$\d[\d,.]*[KMB]?"                       # $30K, $2M, $1,500
    r"|\b\d+(?:\.\d+)?[KMB]\b"                 # 30K, 11M, 2.5B
    r"|\b\d+ ?(?:million|billion|thousand)\b"  # 11 million
    r"|\d+%"                                   # 30%
    r"|\d+\+",                                 # 40+
    re.IGNORECASE,
)


def check_9_numeric_grounding(document_xml, career_file_text):
    """Check 9: every metric in the rendered CV traces to the career file.

    Catches invented/inflated numbers (e.g. a "30%" or "$50K" the career file
    never states). Flags percentages, count claims (40+), currency amounts,
    and K/M/B / million-style magnitudes — and only when the digit sequence
    appears nowhere in the career file (checked against both the raw text and
    a digits-only squash, so "$30,000" grounds "30000"). A real number written
    slightly differently still passes. Semantic inflation ("supported" ->
    "led") is the editorial honesty companion, not this check.

    Skipped when no career file text is provided.
    """
    if not career_file_text:
        return True, "No career file provided; numeric grounding skipped."

    text = _visible_text(document_xml)
    metrics = _METRIC_RE.findall(text)
    career_squashed = re.sub(r"\D", "", career_file_text)
    ungrounded = []
    for m in metrics:
        digits = re.sub(r"\D", "", m)
        if digits and digits not in career_file_text \
                and digits not in career_squashed:
            ungrounded.append(m)
    # de-dup while keeping order
    seen = set()
    ungrounded = [m for m in ungrounded if not (m in seen or seen.add(m))]
    ok = not ungrounded
    if ok:
        note = (str(len(metrics)) + " metric(s) checked; all trace to the "
                "career file.")
    else:
        note = ("Metric(s) with no source in the career file: "
                + ", ".join(ungrounded) + ". A bullet may re-frame a real fact "
                "but may not invent a number. Remove or correct, or add the fact "
                "to the career file if it is real.")
    return ok, note


# Generic filler phrases that signal a weak, un-tailored bullet. Each one
# was an actual offender in the 2026-06-25 Cairo batch, where bullets said
# "enterprise decision-makers" while named proof points (Deloitte, Harvard
# Law Review, W3C) sat unused in the career file. High-precision by design:
# only multi-word abstractions. Bare "stakeholders" is deliberately NOT listed
# (it is a common, valid JD term).
#
# These phrases are flagged only when the bullet carrying them has NO concrete
# proof point of its own (no number, no named entity) — see _has_concrete_proof.
# A grounded bullet may legitimately contain one of these phrases ("managed
# analytical workstreams for 40+ multinationals across Technology and Telecom"
# is strong; the phrase is incidental). The unconditional substring ban was the
# v1.6.0 over-correction that pushed bullets toward thin paraphrase — it failed
# the same phrasing the rich benchmark CVs legitimately use.
WEAK_GENERIC_PHRASES = (
    "enterprise decision-makers",
    "enterprise decision makers",
    "global process owners",
    "analytical workstreams",
    "client-ready",
    "actionable insights",
    "actionable recommendations",
    "evidence-based reports",
)

# ponytail: heuristic proof detector, not an NER. A digit or a mid-sentence
# proper noun is "concrete enough"; editorial check 3 is the real judge.
_PROPER_NOUN_RE = re.compile(r"\b([A-Z][a-zA-Z]+|[A-Z]{2,})\b")

# Capitalized tokens that are NOT proof even when the career file contains
# them: sector nouns, languages, months, and sentence-starter noise. Without
# this, "across Technology and Telecom" would count as grounded because the
# career file names those sectors.
_GENERIC_TOKEN_STOPLIST = frozenset({
    "Technology", "Telecom", "Telecommunications", "Media", "Health",
    "English", "German", "Arabic", "The", "An", "It", "This", "These",
    "January", "February", "March", "April", "May", "June", "July",
    "August", "September", "October", "November", "December", "Present",
})

_CAP_TOKEN_RE = re.compile(r"\b[A-Z][A-Za-z0-9+&.-]+\b")


def _career_whitelist(career_file_text):
    """Capitalized tokens from the career file, minus the generic stoplist.

    These are the named entities a bullet may cite as proof: Deloitte,
    Python, Statista, MENA, W3C, Spotify. Membership is what separates a
    career-grounded proper noun from decorative capitalization.
    """
    if not career_file_text:
        return frozenset()
    tokens = set(_CAP_TOKEN_RE.findall(career_file_text))
    return frozenset(t for t in tokens if t not in _GENERIC_TOKEN_STOPLIST)


def _has_concrete_proof(bullet_text):
    """True if the bullet carries a number or a named entity (no career file).

    Fallback detector when no career file is supplied. A number is
    unambiguous proof. For named entities we strip any leading `Label:`
    lead-in (labeled bullet_style) so the capability label's own
    capitalization does not count, then look for a capitalized token that is
    not the first word of the clause. Crude but high-precision for the job.
    """
    if any(ch.isdigit() for ch in bullet_text):
        return True
    # Strip a short leading "Label:" segment (labeled mode) before noun scan.
    clause = re.sub(r"^[^:]{0,60}:\s*", "", bullet_text).strip()
    if not clause:
        return False
    first_word = clause.split()[0]
    for m in _PROPER_NOUN_RE.finditer(clause):
        if m.group(0) != first_word:  # ignore a capitalized sentence start
            return True
    return False


def _is_proofed(bullet_text, whitelist):
    """True if the bullet carries a digit or a career-file named entity.

    The whitelist version of _has_concrete_proof: a capitalized token counts
    only when the career file actually contains it (and it is not a stoplisted
    sector/language noun), so "Technology and Telecom" no longer grounds a
    bullet while "Deloitte" and "Python" still do.
    """
    if any(ch.isdigit() for ch in bullet_text):
        return True
    clause = re.sub(r"^[^:]{0,60}:\s*", "", bullet_text).strip()
    for m in _CAP_TOKEN_RE.finditer(clause):
        if m.start() == 0:
            continue  # sentence-start capital is not a named entity
        if m.group(0) in whitelist:
            return True
    return False


def check_10_bullet_strength(experiences, career_file_text=None):
    """Check 10: bullet strength — generic-filler test + proof-density floor.

    Two layers:

    1. **Weak-phrase test** (always on): a WEAK_GENERIC_PHRASE fails only
       when its bullet has no concrete proof of its own. A grounded bullet
       may keep natural phrasing.
    2. **Proof-density floor** (when `career_file_text` is supplied): per
       slot, most bullets must be proofed — carry a digit or a capitalized
       token that appears in the career file (minus the generic stoplist).
       Floor: slots with >= 3 bullets need >= 2 proofed; 2-bullet slots need
       >= 1. One interpretive/ungrounded bullet per 3-bullet slot is allowed
       by design — that is the domain-translation pattern, not a defect.

    Without a career file the density floor is skipped and proof detection
    falls back to the heuristic _has_concrete_proof (any mid-clause
    capitalized token) — the pre-v1.8.0 behavior, which the 2026-06-27
    Berlin batch showed is trivially satisfied ("Technology", a year).
    """
    if not experiences:
        return True, "No experiences; bullet-strength check skipped."

    whitelist = _career_whitelist(career_file_text)
    if whitelist:
        def proofed(text):
            return _is_proofed(text, whitelist)
    else:
        proofed = _has_concrete_proof

    hits = []
    density = []
    for i, e in enumerate(experiences):
        bullets = [_bullet_text(b) for b in (e.get("bullets") or [])]
        flags = [proofed(t) for t in bullets]
        for btext, is_proofed_flag in zip(bullets, flags):
            if is_proofed_flag:
                continue  # grounded bullet — phrase is incidental, allowed
            low = btext.lower()
            for phrase in WEAK_GENERIC_PHRASES:
                if phrase in low:
                    hits.append("slot " + str(i + 1) + " ("
                                + str(e.get("company", "?")) + "): '" + phrase
                                + "' in an ungrounded bullet")
        if whitelist and bullets:
            floor = 2 if len(bullets) >= 3 else 1
            n_proofed = sum(flags)
            if n_proofed < floor:
                unproofed = [t[:50] for t, f in zip(bullets, flags) if not f]
                density.append(
                    "slot " + str(i + 1) + " ("
                    + str(e.get("company", "?")) + "): only "
                    + str(n_proofed) + "/" + str(len(bullets))
                    + " bullets carry a digit or career-file named entity "
                    "(floor " + str(floor) + "). Unproofed: " + str(unproofed))

    ok = not hits and not density
    if ok:
        note = ("No ungrounded generic-filler phrasing; proof density "
                + ("met per slot." if whitelist
                   else "heuristic only (no career file supplied)."))
    else:
        note = ("; ".join(hits + density)
                + ". Ground the flagged bullets with the slot's named proof "
                "point / metric from the career file (see the diagnosis's "
                "per-slot proof points), or drop the abstraction.")
    return ok, note


def _parse_slot_proof_points(diagnosis_text):
    """Parse `- Slot N ... proof point: <text>` lines from a Diagnosis.md.

    Returns {slot_number: proof_point_string}. The diagnosis file is the
    source of truth — a model-copied content_map field would be a dodge
    surface (copy a weaker proof point, pass the check).
    """
    points = {}
    for line in diagnosis_text.splitlines():
        line = line.strip()
        m = re.match(r"[-*]\s*Slot\s*(\d+)", line, re.IGNORECASE)
        if not m:
            continue
        pp = re.search(r"proof point:\s*(.+?)(?:\||$)", line, re.IGNORECASE)
        if pp:
            points[int(m.group(1))] = pp.group(1).strip()
    return points


def _distinctive_tokens(proof_point_text):
    """Numbers and career-grade capitalized tokens from a proof-point string.

    These are the strings whose presence in the slot's bullets proves the
    proof point was surfaced. Returns [] for 'none' or a proof point with
    nothing distinctive (the check then skips loudly rather than false-fail).
    """
    text = proof_point_text.strip().rstrip(".")
    if text.lower() in ("none", "n/a", "no proof point"):
        return []
    tokens = re.findall(r"\$?\d[\d,.%+]*[KMB]?", text)
    tokens += [t for t in _CAP_TOKEN_RE.findall(text)
               if t not in _GENERIC_TOKEN_STOPLIST]
    return tokens


def check_11_proof_points(experiences, diagnosis_md_path):
    """Check 11: each slot's diagnosis proof point surfaces in its bullets.

    Check 8 proves a slot carries *a* keyword; this proves the slot carries
    *its assigned proof point* — the named credential/metric the diagnosis
    said the bullets must surface. Catches the observed drift where the lead
    slot's "+30% publication speed" proof point silently fell out of the
    rendered bullets (2026-06-27 Berlin batch).

    A slot passes when >= 1 distinctive token from its proof point appears
    (word-bounded, case-insensitive) in its bullets. Proof points with no
    distinctive token skip loudly. Skipped entirely when no diagnosis file.
    """
    if not diagnosis_md_path or not os.path.exists(diagnosis_md_path):
        return True, "No diagnosis file supplied; proof-point check skipped."
    with open(diagnosis_md_path, encoding="utf-8", errors="replace") as f:
        text = f.read()
    points = _parse_slot_proof_points(text)
    if not points:
        return True, ("No 'Slot N ... proof point:' lines found in the "
                      "diagnosis; check skipped (lint_diagnosis enforces "
                      "their presence upstream).")

    problems, skipped = [], []
    for slot_no, pp in sorted(points.items()):
        idx = slot_no - 1
        if idx >= len(experiences):
            continue
        tokens = _distinctive_tokens(pp)
        if not tokens:
            skipped.append("slot %d (proof point %r has no distinctive "
                           "token)" % (slot_no, pp[:40]))
            continue
        blob = " ".join(
            _bullet_text(b)
            for b in (experiences[idx].get("bullets") or [])).lower()
        found = any(
            re.search(r"(?<!\w)" + re.escape(t.lower()) + r"(?!\w)", blob)
            for t in tokens)
        if not found:
            problems.append(
                "slot %d (%s): none of the diagnosis proof-point tokens %s "
                "appear in its bullets (proof point: %r)"
                % (slot_no, experiences[idx].get("company", "?"),
                   tokens[:6], pp[:60]))

    ok = not problems
    if ok:
        note = "Every slot surfaces its diagnosis proof point."
        if skipped:
            note += " Skipped (no distinctive token): " + "; ".join(skipped)
    else:
        note = ("; ".join(problems)
                + ". The proof point is the slot's assigned credential — "
                "put it back in a bullet, or fix the diagnosis if the "
                "assignment changed.")
        if skipped:
            note += " Skipped: " + "; ".join(skipped)
    return ok, note


def check_12_education_completeness(degrees, document_xml,
                                    expected_degree_count=None):
    """Check 12: every degree renders; none silently dropped.

    The 2026-07-14 Werkstudent CV shipped without the BA because the old
    two-slot education template forced a three-degree candidate to drop one,
    and nothing checked. Now education is a `degrees` loop and this check
    asserts (a) the content_map carries at least `cv.expected_degree_count`
    degrees (the count job-search-setup records from the career file), and
    (b) every degree's institution is visible in the rendered document.

    Without `expected_degree_count` the count cannot be enforced — the note
    says so loudly instead of pretending coverage.
    """
    if not degrees:
        return False, ("content_map has no degrees; education is a required "
                       "section — every degree in the career file renders.")
    problems = []
    if expected_degree_count and len(degrees) < expected_degree_count:
        problems.append(
            "%d degree(s) in content_map; career file has %d "
            "(cv.expected_degree_count). A degree was dropped."
            % (len(degrees), expected_degree_count))
    text = _visible_text(document_xml)
    for i, deg in enumerate(degrees):
        inst = deg.get("institution", "")
        if inst and _norm_ws(inst) not in _norm_ws(text):
            problems.append(
                "degrees[%d] institution %r not visible in the rendered CV"
                % (i, inst))
    ok = not problems
    if ok:
        note = "%d degree(s) present and visible" % len(degrees)
        note += ("." if expected_degree_count else
                 "; cv.expected_degree_count not set, so a dropped degree "
                 "cannot be detected — set it in config.yaml.")
    else:
        note = "; ".join(problems)
    return ok, note


def scan_batch_sameyness(session_dir):
    """Batch-level sweep (WARN only, not part of run_full_audit): exact
    duplicate experience bullets across different CVs in a session folder.

    Cross-CV reuse of a bullet is sometimes legitimate (the same true fact
    for two similar JDs) — the sweep exists so it is a visible choice, not
    silent drift (2026-06-14 Denmark: one slot byte-identical across all
    ten CVs). Returns a list of warning strings, empty when clean.
    """
    bullet_files = {}
    for fname in sorted(os.listdir(session_dir)):
        if not (fname.startswith("CV - ") and fname.endswith(".docx")):
            continue
        doc = Document(os.path.join(session_dir, fname))
        in_experience = False
        for p in doc.paragraphs:
            text = p.text.strip()
            if text == "PROFESSIONAL EXPERIENCE":
                in_experience = True
                continue
            if text == "EDUCATION":
                in_experience = False
            if not in_experience or not text:
                continue
            # Title rows carry a tab; company lines carry the middle dot.
            if "\t" in text or "·" in text:
                continue
            norm = " ".join(text.split()).lower()
            bullet_files.setdefault(norm, set()).add(fname)

    warnings = []
    reported = set()
    for norm, files in sorted(bullet_files.items()):
        if len(files) > 1:
            warnings.append("bullet shared by %s: %r"
                            % (", ".join(sorted(files)), norm[:70]))
            reported.add(norm)
    # Clause-level pass: same clause behind different labels.
    clause_files = {}
    for norm, files in bullet_files.items():
        clause = re.sub(r"^[^:]{0,60}:\s*", "", norm)
        if clause != norm and clause:
            clause_files.setdefault(clause, set()).update(files)
    for clause, files in sorted(clause_files.items()):
        if len(files) > 1 and clause not in reported:
            warnings.append("clause shared (different labels) by %s: %r"
                            % (", ".join(sorted(files)), clause[:70]))
    return warnings


def run_full_audit(rendered_docx_path, diagnosis_md_path, content_map,
                   expected_keywords, expect_bold=True, career_file_path=None,
                   bold_plan=None, require_editorial=True,
                   expected_degree_count=None):
    """Run the programmatic audit checks. Returns an AuditResult.

    `bold_plan` is the plan returned by build_bold_plan(); Check 5 uses it
    to verify each planned bold span rendered as a real bold run.

    The editorial checks (1: lead slots serve the diagnosed problem with
    their proof points; 3: recruiter-fit / richness / domain translation /
    honesty companion) are seeded as FAILED when `require_editorial` is True
    (the default). `all_passed` stays False until the model records a
    verdict for each via result.record_editorial(name, ok, note) — a CV can
    no longer pass by omission. Recording the verdicts is authoring work,
    not a pause.
    """
    document_xml = _read_document_xml(rendered_docx_path)
    experiences = content_map.get("experiences", [])
    result = AuditResult()

    if require_editorial:
        result.passed["check_1_lead_slots"] = False
        result.notes["check_1_lead_slots"] = (
            "editorial verdict not recorded. Read the lead slots against the "
            "diagnosis's problem statement and per-slot proof points, then "
            "call result.record_editorial('check_1_lead_slots', ok, note).")
        result.passed["check_3_recruiter_fit"] = False
        result.notes["check_3_recruiter_fit"] = (
            "editorial verdict not recorded. Judge richness vs the career "
            "file, domain translation, recruiter-fit, and the honesty "
            "companion (no semantic inflation), then call "
            "result.record_editorial('check_3_recruiter_fit', ok, note).")

    ok2, note2 = check_2_keywords_in_experience(experiences, expected_keywords)
    result.passed["check_2_keywords"] = ok2
    result.notes["check_2_keywords"] = note2

    ok4, note4 = check_4_ampersands(document_xml, content_map)
    result.passed["check_4_ampersands"] = ok4
    result.notes["check_4_ampersands"] = note4

    ok5, note5 = check_5_rendered_integrity(
        rendered_docx_path, content_map,
        bold_plan=bold_plan, expect_bold=expect_bold)
    result.passed["check_5_integrity"] = ok5
    result.notes["check_5_integrity"] = note5

    ok6, note6 = check_6_no_em_dashes(document_xml)
    result.passed["check_6_em_dashes"] = ok6
    result.notes["check_6_em_dashes"] = note6

    ok7, note7 = check_7_experience_structure(experiences)
    result.passed["check_7_structure"] = ok7
    result.notes["check_7_structure"] = note7

    ok8, note8 = check_8_slot_coverage(experiences, expected_keywords)
    result.passed["check_8_coverage"] = ok8
    result.notes["check_8_coverage"] = note8

    career_text = None
    if career_file_path:
        with open(career_file_path, encoding="utf-8", errors="replace") as f:
            career_text = f.read()
    ok9, note9 = check_9_numeric_grounding(document_xml, career_text)
    result.passed["check_9_grounding"] = ok9
    result.notes["check_9_grounding"] = note9

    ok10, note10 = check_10_bullet_strength(experiences, career_text)
    result.passed["check_10_bullet_strength"] = ok10
    result.notes["check_10_bullet_strength"] = note10

    ok11, note11 = check_11_proof_points(experiences, diagnosis_md_path)
    result.passed["check_11_proof_points"] = ok11
    result.notes["check_11_proof_points"] = note11

    ok12, note12 = check_12_education_completeness(
        content_map.get("degrees", []), document_xml,
        expected_degree_count=expected_degree_count)
    result.passed["check_12_education"] = ok12
    result.notes["check_12_education"] = note12

    return result


def _selftest():
    """Smallest check that fails if check_10 logic breaks (weak/strong split)."""
    weak = [{"company": "Statista", "bullets": [
        "Tracked competitive positioning for enterprise decision-makers."]}]
    strong = [{"company": "Statista", "bullets": [
        "Synthesized findings into reports cited by Deloitte and the "
        "Harvard Law Review, briefing global stakeholders."]}]
    ok_weak, _ = check_10_bullet_strength(weak)
    ok_strong, _ = check_10_bullet_strength(strong)
    assert not ok_weak, "check_10 should FAIL a generic-filler bullet"
    assert ok_strong, "check_10 should PASS a named-proof-point bullet"
    # Grounding-aware: a generic phrase is allowed when the bullet itself
    # carries concrete proof (a number or named entity).
    grounded = [{"company": "Statista", "bullets": [
        "Managed analytical workstreams for 40+ multinationals across "
        "Technology and Telecom, delivering client-ready outputs."]}]
    ok_grounded, _ = check_10_bullet_strength(grounded)
    assert ok_grounded, ("check_10 should PASS a grounded bullet even when it "
                         "contains a phrase from WEAK_GENERIC_PHRASES")
    assert _has_concrete_proof("Managed workstreams for 40+ corporations")
    assert not _has_concrete_proof("served enterprise decision-makers")
    # Works on labeled-style bullets too (bold label lead-in, plain string —
    # since v1.8.0 bullets are always plain strings by the time checks run).
    labeled_weak = [{"company": "X", "bullets": [
        "Coverage: served enterprise decision-makers."]}]
    ok_labeled, _ = check_10_bullet_strength(labeled_weak)
    assert not ok_labeled, "check_10 should FAIL filler behind a label lead-in"
    # Whitelist mode: sector nouns do not ground a bullet; career entities do.
    career = ("Synthesized reports cited by Deloitte and W3C. "
              "Built a Python pipeline, +30% speed, across Technology "
              "and Telecom sectors.")
    sector_only = [{"company": "X", "bullets": [
        "Coverage: tracked positioning across Technology and Telecom.",
        "Coverage two: tracked more positioning across sectors broadly.",
        "Reporting: synthesized findings cited by Deloitte and W3C.",
    ]}]
    ok_density, note_density = check_10_bullet_strength(sector_only, career)
    assert not ok_density, "density floor should FAIL 1/3 proofed"
    two_proofed = [{"company": "X", "bullets": [
        "Coverage: tracked positioning for 40+ multinationals.",
        "Coverage two: tracked more positioning across sectors broadly.",
        "Reporting: synthesized findings cited by Deloitte and W3C.",
    ]}]
    ok_two, _ = check_10_bullet_strength(two_proofed, career)
    assert ok_two, "2/3 proofed should PASS the density floor"
    print("audit self-test (check_10): passed.")


if __name__ == "__main__":
    import sys
    if len(sys.argv) >= 2 and sys.argv[1] == "--selftest":
        _selftest()
        sys.exit(0)
    if len(sys.argv) >= 3 and sys.argv[1] == "--sameyness":
        warnings = scan_batch_sameyness(sys.argv[2])
        if warnings:
            print("Sameyness sweep: %d duplicate(s) across CVs" % len(warnings))
            for w in warnings:
                print("  WARN:", w)
        else:
            print("Sameyness sweep: clean (no duplicate bullets across CVs).")
        sys.exit(0)
    if len(sys.argv) < 2:
        print("Usage: python audit.py <rendered_cv.docx>"
              "  |  python audit.py --selftest"
              "  |  python audit.py --sameyness <session_dir>")
        sys.exit(1)
    if not os.path.isfile(sys.argv[1]):
        print(f"Error: file not found: {sys.argv[1]}", file=sys.stderr)
        sys.exit(1)
    path = sys.argv[1]
    xml = _read_document_xml(path)
    doc = Document(path)
    paras = doc.paragraphs
    empty = sum(1 for p in paras if not p.text.strip())
    with zipfile.ZipFile(path) as z:
        rels = z.read("word/_rels/document.xml.rels").decode(
            "utf-8", errors="replace")
    print("Paragraphs (python-docx):", len(paras))
    print("Empty paragraphs:", empty,
          " <- empty bullets = the 2026-05-11/test5 corruption; investigate")
    print("Escaped ampersands:", xml.count("&amp;"))
    print("Hyperlink relationships:", rels.count("/relationships/hyperlink"))
    print("Em dashes:", _visible_text(xml).count("—"))
