"""
Tests for the v1.8.0 render pipeline: plain-string render + postprocess bold
+ section removal. This is the architecture that makes the 2026-05-11 /
test5 corruption (RichText embedded inside <w:t> -> invisible bullets)
structurally impossible: bullets are plain strings at render time, and bold
is applied afterwards by cloning the rendered run.
"""
import zipfile

import pytest
from docx import Document
from docxtpl import DocxTemplate

from md_to_richtext import build_bold_plan
from postprocess import postprocess_cv, PostprocessError, _split_run_bold
from render_cv import render, effective_sections
from conftest import TEMPLATE, REPO_ROOT, contact_links, minimal_content_map


def _render_plain(cm, out_path):
    """Render a content map (already marker-stripped) through the template."""
    tpl = DocxTemplate(TEMPLATE)
    tpl.render(cm, autoescape=True)
    tpl.save(str(out_path))
    return str(out_path)


def _para_with_text(doc, text):
    norm = " ".join(text.split())
    for p in doc.paragraphs:
        if " ".join(p.text.split()) == norm:
            return p
    raise AssertionError("paragraph not found: %r" % text[:60])


def test_labeled_bullet_bold_label_formatting_preserved(tmp_path):
    bullet = "**Pipeline automation:** built a Python pipeline cutting time 30%."
    cm = minimal_content_map(experiences=[{
        "title": "Senior Analyst", "dates": "2023-Present",
        "company": "Acme", "location": "City", "bullets": [bullet],
    }])
    cm, plan = build_bold_plan(cm, mode="labeled")
    path = _render_plain(cm, tmp_path / "labeled.docx")
    summary = postprocess_cv(path, plan)
    assert summary["bolded_runs"] == 1

    doc = Document(path)
    p = _para_with_text(
        doc, "Pipeline automation: built a Python pipeline cutting time 30%.")
    assert p.runs[0].text == "Pipeline automation:"
    assert p.runs[0].bold is True
    assert p.runs[1].bold is not True
    # Cloned runs inherit the template bullet run's explicit formatting.
    assert p.runs[0].font.name == "Calibri"
    assert p.runs[0].font.size.pt == 10.0
    assert p.runs[1].font.name == "Calibri"


def test_inline_two_spans_alternation(tmp_path):
    bullet = "Lifted **activation by 18%**, covered in **TechCrunch**."
    cm = minimal_content_map(experiences=[{
        "title": "T", "dates": "D", "company": "Acme", "location": "L",
        "bullets": [bullet],
    }])
    cm, plan = build_bold_plan(cm, mode="inline")
    path = _render_plain(cm, tmp_path / "inline.docx")
    postprocess_cv(path, plan)

    doc = Document(path)
    p = _para_with_text(
        doc, "Lifted activation by 18%, covered in TechCrunch.")
    texts = [r.text for r in p.runs]
    bolds = [bool(r.bold) for r in p.runs]
    assert texts == ["Lifted ", "activation by 18%", ", covered in ",
                     "TechCrunch", "."]
    assert bolds == [False, True, False, True, False]


def test_bold_and_ampersand_interaction(tmp_path):
    bullet = "**A & B:** rest & more."
    cm = minimal_content_map(experiences=[{
        "title": "T", "dates": "D", "company": "Acme", "location": "L",
        "bullets": [bullet],
    }])
    cm, plan = build_bold_plan(cm, mode="labeled")
    path = _render_plain(cm, tmp_path / "amp.docx")
    postprocess_cv(path, plan)

    doc = Document(path)
    p = _para_with_text(doc, "A & B: rest & more.")
    assert p.runs[0].bold is True
    assert p.runs[0].text == "A & B:"
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert "A &amp; B:" in xml  # ampersand escaped, not stripped


def test_split_survives_pre_split_paragraph(tmp_path):
    """Defensive: a bullet paragraph that arrives as multiple runs is merged
    before splitting, so text and spans stay correct."""
    import copy as _copy
    bullet = "**Label:** the rest of it."
    cm = minimal_content_map(experiences=[{
        "title": "T", "dates": "D", "company": "Acme", "location": "L",
        "bullets": [bullet],
    }])
    cm, plan = build_bold_plan(cm, mode="labeled")
    path = _render_plain(cm, tmp_path / "presplit.docx")

    # Artificially split the bullet's single run into two.
    doc = Document(path)
    p = _para_with_text(doc, "Label: the rest of it.")
    src = p.runs[0]
    full = src.text
    clone = _copy.deepcopy(src._r)
    src._r.addnext(clone)
    src.text = full[:7]
    from docx.text.run import Run
    Run(clone, p).text = full[7:]
    assert len(p.runs) == 2

    _split_run_bold(p, plan[0]["spans"])
    assert " ".join(p.text.split()) == "Label: the rest of it."
    assert p.runs[0].text == "Label:"
    assert p.runs[0].bold is True


def test_section_removal_summary_and_additional(tmp_path):
    cm = minimal_content_map()
    cm, plan = build_bold_plan(cm, mode="plain")
    path = _render_plain(cm, tmp_path / "sections.docx")
    summary = postprocess_cv(path, plan,
                             disabled_sections=("summary", "additional"))
    assert set(summary["removed_sections"]) == {"summary", "additional"}

    doc = Document(path)
    texts = [p.text.strip() for p in doc.paragraphs]
    assert "PROFESSIONAL SUMMARY" not in texts
    assert "ADDITIONAL" not in texts
    # Neighboring sections intact.
    assert "CORE SKILLS" in texts
    assert "EDUCATION" in texts
    assert any("Did a plain thing." in t for t in texts)


def test_required_sections_never_removed(tmp_path):
    cm = minimal_content_map()
    cm, plan = build_bold_plan(cm, mode="plain")
    path = _render_plain(cm, tmp_path / "required.docx")
    summary = postprocess_cv(path, plan, disabled_sections=("experience",))
    assert summary["removed_sections"] == []
    assert any("required" in w for w in summary["warnings"])
    doc = Document(path)
    assert "PROFESSIONAL EXPERIENCE" in [p.text.strip() for p in doc.paragraphs]


def test_hyperlinks_point_at_their_labels_after_roundtrip(tmp_path):
    """Section removal must not orphan or misdirect the contact links.

    Asserting a COUNT here (as this test used to) is what let the 2026-08
    leak through: both rels were present and both pointed at the template
    author. Assert the pairing instead.
    """
    cm = minimal_content_map()
    cm, plan = build_bold_plan(cm, mode="plain")
    path = _render_plain(cm, tmp_path / "links.docx")
    postprocess_cv(path, plan, disabled_sections=("summary",),
                   contact_links=contact_links(cm))
    from audit import _read_hyperlinks, _norm_target
    links = {label: target for label, _rid, target in _read_hyperlinks(path)}
    for key in ("personal_site", "linkedin_url"):
        assert _norm_target(links[cm[key]]) == _norm_target(cm[key])


def test_plan_mismatch_raises(tmp_path):
    cm = minimal_content_map()
    cm, plan = build_bold_plan(cm, mode="plain")
    path = _render_plain(cm, tmp_path / "mismatch.docx")
    bad_plan = [{"section": "experience", "slot": 0,
                 "text": "a bullet that was never rendered", "spans": []}]
    with pytest.raises(PostprocessError):
        postprocess_cv(path, bad_plan)


# ---------------------------------------------------------------------------
# End-to-end through render(): validation -> render -> postprocess -> audit.
# ---------------------------------------------------------------------------

def _three_slot_map():
    # Bullet counts meet the v1.9.0 floors (lead >= 5, slot 2 >= 4,
    # slot 3 >= 3) and the v2.0.0 length floors (clause >= 12 words,
    # section mean >= 20) — career-file density is the contract, and the
    # fixture models the standard a real content_map must meet.
    return minimal_content_map(experiences=[
        {"title": "Senior Analyst", "dates": "2023 - Present",
         "company": "Acme", "location": "City", "end_year": 9999,
         "bullets": [
             "**Pipeline automation:** diagnosed a manual data entry "
             "bottleneck and built an automated Python pipeline that "
             "increased report publication speed by 30% across weekly "
             "reporting cycles.",
             "**Dashboard ownership:** owned Power BI KPI dashboards used "
             "across four departments, gathering stakeholder requirements "
             "and shipping monthly iterations consumed in weekly executive "
             "reporting.",
             "**Executive interviews:** conducted structured interviews "
             "with senior executives and subject-matter experts, "
             "translating qualitative input into decision evidence used by "
             "enterprise clients for strategic planning decisions.",
             "**Market scoping:** scoped ambiguous research questions into "
             "repeatable methodologies covering 40+ multinational "
             "corporations, delivering competitive landscape reports under "
             "tight publication deadlines.",
             "**Citation impact:** produced analyses cited by Alpha Advisory and "
             "W3C in published industry research, a quality signal "
             "recognized by institutional clients worldwide.",
         ]},
        {"title": "Analyst", "dates": "2020 - 2023",
         "company": "Acme", "location": "City", "end_year": 2023,
         "bullets": [
             "**Data quality:** cleaned and validated epidemiological and "
             "market data across high-frequency publication cycles, "
             "maintaining accuracy across 50+ regional datasets during "
             "peak demand.",
             "**Trial tracking:** tracked vaccine trials from Phase I to "
             "Phase III, structuring findings into reports used by "
             "institutional and government clients worldwide during a "
             "period of unprecedented global demand.",
             "**Workflow automation:** supported Python data workflows "
             "that reduced manual entry across the team, freeing analyst "
             "time for higher-value interpretive work across high-frequency "
             "publication cycles.",
             "**Progression:** promoted from assistant to expert within 3 "
             "years, building fluency in hypothesis structuring, source "
             "triangulation, and executive-audience writing.",
         ]},
        {"title": "Consultant", "dates": "2018 - 2020",
         "company": "Other Co", "location": "Town", "end_year": 2020,
         "bullets": [
             "**Primary research:** conducted structured expert interviews "
             "for global market research projects, ensuring accurate "
             "domain capture across specialized industries for "
             "international clients.",
             "**Cross-cultural delivery:** managed bilingual communication "
             "and coordination for research projects targeting regional "
             "markets, delivering outputs aligned with international "
             "research standards.",
             "**Client synthesis:** structured qualitative findings into "
             "deliverables for 20+ client engagements across the region, "
             "consumed directly by strategy and product teams for planning "
             "and prioritization.",
         ]},
    ])


KEYWORDS = ["pipeline", "dashboards", "data quality", "market research"]


def test_render_end_to_end_labeled(tmp_path):
    config = {"cv": {"bullet_style": "labeled", "max_experience_slots": 3}}
    out = tmp_path / "cv.docx"
    result = render(
        diagnosis_path=None,
        content_map=_three_slot_map(),
        config=config,
        repo_root=REPO_ROOT,
        output_path=str(out),
        expected_keywords=KEYWORDS,
    )
    assert result.all_passed, result.failure_summary

    # Every bullet must be readable and the labels bold — what Word and an
    # ATS parser see, not what a raw-XML regex sees.
    doc = Document(str(out))
    p = _para_with_text(
        doc, "Pipeline automation: diagnosed a manual data entry bottleneck "
             "and built an automated Python pipeline that increased report "
             "publication speed by 30% across weekly reporting cycles.")
    assert p.runs[0].bold is True
    assert p.runs[0].text == "Pipeline automation:"


def test_render_region_override_removes_summary(tmp_path):
    config = {"cv": {
        "bullet_style": "labeled",
        "max_experience_slots": 3,
        "region_section_overrides": {"EU": {"summary": False}},
    }}
    cm = _three_slot_map()
    del cm["summary"]  # not required when the region disables the section
    out = tmp_path / "cv_eu.docx"
    result = render(
        diagnosis_path=None,
        content_map=cm,
        config=config,
        repo_root=REPO_ROOT,
        output_path=str(out),
        expected_keywords=KEYWORDS,
        region="EU",
    )
    assert result.all_passed, result.failure_summary
    texts = [p.text.strip() for p in Document(str(out)).paragraphs]
    assert "PROFESSIONAL SUMMARY" not in texts


def test_effective_sections_resolution():
    config = {"cv": {"region_section_overrides": {"EU": {"summary": False}}}}
    enabled, disabled = effective_sections(config, region="EU")
    assert "summary" in disabled and "summary" not in enabled
    enabled_us, disabled_us = effective_sections(config, region="US")
    assert disabled_us == [] and "summary" in enabled_us
    enabled_none, disabled_none = effective_sections(config, region=None)
    assert disabled_none == []


def test_validate_rejects_missing_end_year_and_thin_slots(tmp_path):
    from render_cv import validate_content_map
    config = {"cv": {"max_experience_slots": 3}}
    cm = _three_slot_map()
    del cm["experiences"][0]["end_year"]
    cm["experiences"][1]["bullets"] = cm["experiences"][1]["bullets"][:1]
    with pytest.raises(ValueError) as exc:
        validate_content_map(cm, config)
    msg = str(exc.value)
    assert "end_year" in msg
    assert "floor" in msg


def test_validate_rejects_unlabeled_bullet_in_labeled_mode():
    from render_cv import validate_content_map
    config = {"cv": {"max_experience_slots": 3}}
    cm = _three_slot_map()
    cm["experiences"][0]["bullets"][0] = "A bullet without any label lead-in."
    with pytest.raises(ValueError) as exc:
        validate_content_map(cm, config, mode="labeled")
    assert "labeled mode" in str(exc.value)


def test_validate_rejects_em_dash():
    from render_cv import validate_content_map
    config = {"cv": {"max_experience_slots": 3}}
    cm = _three_slot_map()
    cm["summary"] = "A summary with an em dash — banned."
    with pytest.raises(ValueError) as exc:
        validate_content_map(cm, config)
    assert "em dash" in str(exc.value)


def test_validate_default_floors_are_5_4_3():
    """The 2026-07-14 Werkstudent regression: 4/3/3 must now fail."""
    from render_cv import validate_content_map
    config = {"cv": {"max_experience_slots": 3}}
    cm = _three_slot_map()
    cm["experiences"][0]["bullets"] = cm["experiences"][0]["bullets"][:4]
    cm["experiences"][1]["bullets"] = cm["experiences"][1]["bullets"][:3]
    with pytest.raises(ValueError) as exc:
        validate_content_map(cm, config)
    msg = str(exc.value)
    assert "floor is 5" in msg
    assert "floor is 4" in msg


def test_validate_rejects_thin_written_bullets():
    """v2.0.0: the thin-CV leak — count floors met, bullets written short.

    A 12-word-fragment bullet and a section written at ~16 words/clause both
    fail validation now; before, they passed every gate (the v1.6.0-era
    'category-noun' CVs and the showcase's one-liner bullets)."""
    from render_cv import validate_content_map
    config = {"cv": {"max_experience_slots": 3}}

    # One fragment among rich bullets: per-bullet floor fires.
    cm = _three_slot_map()
    cm["experiences"][0]["bullets"][0] = \
        "**Pipeline automation:** built a Python pipeline, 30% faster."
    with pytest.raises(ValueError) as exc:
        validate_content_map(cm, config, mode="labeled")
    assert "fragment" in str(exc.value)

    # Uniformly short bullets clear the per-bullet floor but fail the mean.
    cm = _three_slot_map()
    for role in cm["experiences"]:
        role["bullets"] = [
            "**Label %d:** delivered analytical outputs across markets for "
            "enterprise clients under sustained deadline pressure "
            "conditions." % i
            for i in range(len(role["bullets"]))]
    with pytest.raises(ValueError) as exc:
        validate_content_map(cm, config, mode="labeled")
    assert "written thin" in str(exc.value)

    # The full fixture passes both length floors.
    validate_content_map(_three_slot_map(), config, mode="labeled")


def test_validate_bullet_floors_config_override():
    from render_cv import validate_content_map
    config = {"cv": {"max_experience_slots": 3, "bullet_floors": [3, 2, 2]}}
    cm = _three_slot_map()
    cm["experiences"][0]["bullets"] = cm["experiences"][0]["bullets"][:3]
    cm["experiences"][1]["bullets"] = cm["experiences"][1]["bullets"][:2]
    validate_content_map(cm, config)  # passes under the explicit override


def test_validate_rejects_retired_msc_ba_keys():
    from render_cv import validate_content_map
    config = {"cv": {"max_experience_slots": 3}}
    cm = _three_slot_map()
    cm["msc_degree"] = "MSc"
    with pytest.raises(ValueError) as exc:
        validate_content_map(cm, config)
    assert "retired" in str(exc.value)


def test_validate_requires_degrees_with_fields_and_bullets():
    from render_cv import validate_content_map
    config = {"cv": {"max_experience_slots": 3}}
    cm = _three_slot_map()
    cm["degrees"] = [{"name": "MSc", "date": "2018",
                      "institution": "A University", "location": "City",
                      "bullets": []}]
    with pytest.raises(ValueError) as exc:
        validate_content_map(cm, config)
    assert "no bullets" in str(exc.value)
    del cm["degrees"]
    with pytest.raises(ValueError) as exc:
        validate_content_map(cm, config)
    assert "degrees" in str(exc.value)


def test_check_12_education_completeness(tmp_path):
    from audit import check_12_education_completeness, _read_document_xml
    cm = _three_slot_map()
    cm2, plan = build_bold_plan(cm, mode="plain")
    path = _render_plain(cm2, tmp_path / "edu.docx")
    xml = _read_document_xml(path)
    degrees = cm2["degrees"]
    ok, note = check_12_education_completeness(degrees, xml,
                                               expected_degree_count=2)
    assert ok, note
    # A dropped degree fails the count assertion.
    ok, note = check_12_education_completeness(degrees[:1], xml,
                                               expected_degree_count=2)
    assert not ok and "dropped" in note
    # A degree in the map but not in the document fails visibility.
    ghost = degrees + [{"name": "PhD", "institution": "Ghost Institute",
                        "bullets": ["x"]}]
    ok, note = check_12_education_completeness(ghost, xml,
                                               expected_degree_count=3)
    assert not ok and "Ghost Institute" in note
    # Unset expected count: passes on visibility, says so loudly.
    ok, note = check_12_education_completeness(degrees, xml)
    assert ok and "expected_degree_count" in note


def test_render_all_degrees_visible_end_to_end(tmp_path):
    """Three degrees through the template loop; none dropped."""
    config = {"cv": {"bullet_style": "labeled", "max_experience_slots": 3,
                     "expected_degree_count": 3}}
    cm = _three_slot_map()
    cm["degrees"] = cm["degrees"] + [{
        "name": "MSc Data Science", "date": "Expected 2027",
        "institution": "C Institute", "location": "City",
        "bullets": ["A third degree bullet."]}]
    # most recent first: in-progress degree leads
    cm["degrees"].sort(key=lambda d: d["date"] == "Expected 2027",
                       reverse=True)
    out = tmp_path / "cv3deg.docx"
    result = render(
        diagnosis_path=None, content_map=cm, config=config,
        repo_root=REPO_ROOT, output_path=str(out),
        expected_keywords=KEYWORDS,
    )
    assert result.all_passed, result.failure_summary
    texts = [p.text for p in Document(str(out)).paragraphs]
    for inst in ("A University", "B College", "C Institute"):
        assert any(inst in t for t in texts), inst


def test_render_student_mode_education_first(tmp_path):
    """student_mode moves EDUCATION above PROFESSIONAL EXPERIENCE while the
    full audit still passes and hyperlinks survive."""
    config = {"cv": {"bullet_style": "labeled", "max_experience_slots": 3,
                     "student_mode": True}}
    out = tmp_path / "cv_student.docx"
    result = render(
        diagnosis_path=None,
        content_map=_three_slot_map(),
        config=config,
        repo_root=REPO_ROOT,
        output_path=str(out),
        expected_keywords=KEYWORDS,
    )
    assert result.all_passed, result.failure_summary

    texts = [p.text.strip() for p in Document(str(out)).paragraphs]
    assert texts.index("EDUCATION") < texts.index("PROFESSIONAL EXPERIENCE")
    # Education content moved with its header; bullets stay readable.
    assert any("A University" in t for t in texts)
    p = _para_with_text(
        Document(str(out)),
        "Pipeline automation: diagnosed a manual data entry bottleneck "
        "and built an automated Python pipeline that increased report "
        "publication speed by 30% across weekly reporting cycles.")
    assert p.runs[0].bold is True
    with zipfile.ZipFile(str(out)) as z:
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
    assert rels.count("/relationships/hyperlink") >= 2


def test_move_section_missing_header_warns(tmp_path):
    """No-op + warning when a header is absent (summary-only doc edge)."""
    cm = minimal_content_map()
    cm, plan = build_bold_plan(cm, mode="plain")
    path = _render_plain(cm, tmp_path / "warn.docx")
    headers = {"experience": "PROFESSIONAL EXPERIENCE",
               "education": "NO SUCH HEADER"}
    summary = postprocess_cv(path, plan, section_headers=headers,
                             student_mode=True)
    assert any("not found" in w for w in summary["warnings"])


def test_transition_mode_allows_one_extra_slot():
    from render_cv import validate_content_map
    config = {"cv": {"max_experience_slots": 3}}
    cm = _three_slot_map()
    cm["experiences"].append({
        "title": "Support Agent", "dates": "2016 - 2018",
        "company": "Old Co", "location": "Town", "end_year": 2018,
        "bullets": [
            "**Client escalations:** resolved technical escalations for "
            "enterprise accounts across three regions, coordinating "
            "engineering and account teams to protect renewal-stage "
            "relationships.",
            "**Onboarding:** trained new team members on the support "
            "playbook and CRM workflows, cutting ramp time for new hires "
            "during a peak growth quarter.",
            "**Knowledge base:** documented recurring issues into a "
            "searchable playbook adopted by the team, reducing repeat "
            "escalations and duplicate investigation work.",
        ],
    })
    # 4 slots rejected in default/adjacent positioning...
    with pytest.raises(ValueError):
        validate_content_map(cm, config, mode="labeled")
    with pytest.raises(ValueError):
        validate_content_map(cm, config, mode="labeled",
                             positioning_mode="adjacent")
    # ...but allowed (max+1) in transition mode.
    validate_content_map(cm, config, mode="labeled",
                         positioning_mode="transition")
